# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = r"C:\Mes Sites Web\Projet_Kiné_Entreprise\M-CHQ_Fonctionnalites.docx"

# ── COLORS ──────────────────────────────────────────────────────────────
BLUE_DARK  = RGBColor(0x1F, 0x38, 0x64)
BLUE_MED   = RGBColor(0x2E, 0x75, 0xB6)
BLUE_LIGHT = RGBColor(0xD9, 0xE2, 0xF3)
RED        = RGBColor(0xC0, 0x00, 0x00)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREY_CELL  = RGBColor(0xF2, 0xF7, 0xFF)

FONT_NAME = "Times New Roman"

# ── HELPERS ─────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def add_header_footer(doc):
    from docx.oxml import OxmlElement
    section = doc.sections[0]

    # Header
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    run1 = p.add_run("AFRILAND FIRST BANK SOUTH SUDAN")
    run1.font.name = FONT_NAME; run1.font.size = Pt(9)
    run1.font.bold = True; run1.font.color.rgb = BLUE_DARK
    run2 = p.add_run("     |     MODULE M-CHQ — Demande de Chéquier en Ligne     |     Version 1.0")
    run2.font.name = FONT_NAME; run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom); pPr.append(pBdr)

    # Footer with page numbers
    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear(); fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fpPr = fp._p.get_or_add_pPr()
    fpBdr = OxmlElement('w:pBdr')
    ftop = OxmlElement('w:top')
    ftop.set(qn('w:val'), 'single'); ftop.set(qn('w:sz'), '6')
    ftop.set(qn('w:space'), '1'); ftop.set(qn('w:color'), '2E75B6')
    fpBdr.append(ftop); fpPr.append(fpBdr)
    r_pg = fp.add_run("Page "); r_pg.font.name = FONT_NAME; r_pg.font.size = Pt(9)
    r_pg.font.color.rgb = RGBColor(0x59,0x59,0x59)
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run_el = OxmlElement('w:r')
    run_el.append(fldChar1); run_el.append(instrText); run_el.append(fldChar2)
    fp._p.append(run_el)
    r_sep = fp.add_run(" / "); r_sep.font.name = FONT_NAME; r_sep.font.size = Pt(9)
    r_sep.font.color.rgb = RGBColor(0x59,0x59,0x59)
    fldChar3 = OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText'); instrText2.text = 'NUMPAGES'
    fldChar4 = OxmlElement('w:fldChar'); fldChar4.set(qn('w:fldCharType'), 'end')
    run_el2 = OxmlElement('w:r')
    run_el2.append(fldChar3); run_el2.append(instrText2); run_el2.append(fldChar4)
    fp._p.append(run_el2)

def configure_styles(doc):
    styles = doc.styles
    # Normal
    n = styles['Normal']
    n.font.name = FONT_NAME; n.font.size = Pt(12)
    nf = n._element.get_or_add_pPr()

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = FONT_NAME; h1.font.size = Pt(16)
    h1.font.bold = True; h1.font.color.rgb = BLUE_DARK
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after  = Pt(12)
    h1.paragraph_format.page_break_before = True

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = FONT_NAME; h2.font.size = Pt(14)
    h2.font.bold = True; h2.font.color.rgb = BLUE_DARK
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after  = Pt(9)
    h2.paragraph_format.page_break_before = False

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = FONT_NAME; h3.font.size = Pt(13)
    h3.font.bold = True; h3.font.color.rgb = BLUE_MED
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after  = Pt(6)

    # List Bullet
    try:
        lb = styles['List Bullet']
        lb.font.name = FONT_NAME; lb.font.size = Pt(12)
        lb.paragraph_format.space_before = Pt(2)
        lb.paragraph_format.space_after  = Pt(2)
    except: pass

    # List Number
    try:
        ln = styles['List Number']
        ln.font.name = FONT_NAME; ln.font.size = Pt(12)
        ln.paragraph_format.space_before = Pt(2)
        ln.paragraph_format.space_after  = Pt(2)
    except: pass

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = FONT_NAME; run.font.size = Pt(16)
        run.font.bold = True; run.font.color.rgb = BLUE_DARK
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = FONT_NAME; run.font.size = Pt(14)
        run.font.bold = True; run.font.color.rgb = BLUE_DARK
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.name = FONT_NAME; run.font.size = Pt(13)
        run.font.bold = True; run.font.color.rgb = BLUE_MED
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(13.8)
    for run in p.runs:
        run.font.name = FONT_NAME; run.font.size = Pt(12)
    return p

def note_critical(doc, label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(13.8)
    r1 = p.add_run(label + " : ")
    r1.font.name = FONT_NAME; r1.font.size = Pt(12)
    r1.font.bold = True; r1.font.color.rgb = RED
    r2 = p.add_run(text)
    r2.font.name = FONT_NAME; r2.font.size = Pt(12)
    r2.font.color.rgb = RED
    return p

def bul(doc, text, level=0):
    style_name = 'List Bullet 2' if level == 1 else 'List Bullet'
    try:
        p = doc.add_paragraph(text, style=style_name)
    except:
        p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(13.8)
    if level == 1:
        p.paragraph_format.left_indent = Cm(1.5)
    for run in p.runs:
        run.font.name = FONT_NAME; run.font.size = Pt(12)
    return p

def num(doc, text):
    try:
        p = doc.add_paragraph(text, style='List Number')
    except:
        p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.line_spacing = Pt(13.8)
    for run in p.runs:
        run.font.name = FONT_NAME; run.font.size = Pt(12)
    return p

def empty(doc):
    p = doc.add_paragraph("")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    return p

def make_table(doc, headers, rows, col_widths_cm):
    total_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=total_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        set_cell_bg(cell, "1F3864")
        set_cell_borders(cell, "2E75B6")
        set_col_width(cell, col_widths_cm[i])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(h)
        run.font.name = FONT_NAME; run.font.size = Pt(10)
        run.font.bold = True; run.font.color.rgb = WHITE

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        fill = "EBF2FF" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            set_cell_bg(cell, fill)
            set_cell_borders(cell, "CCCCCC")
            set_col_width(cell, col_widths_cm[c_idx])
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(cell_text)
            run.font.name = FONT_NAME; run.font.size = Pt(10)

    return tbl

# ════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════════════
doc = Document()

# Page setup: A4, 2cm margins
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2)
section.right_margin  = Cm(2)

configure_styles(doc)
add_header_footer(doc)

# ── SECTION TITLE ──────────────────────────────────────────────────────
p_title = doc.add_paragraph()
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after  = Pt(18)
r = p_title.add_run("V. Détail des Fonctionnalités")
r.font.name = FONT_NAME; r.font.size = Pt(20)
r.font.bold = True; r.font.color.rgb = BLUE_DARK

# ════════════════════════════════════════════════════════════════════════
# FONCTIONNALITÉ 1
# ════════════════════════════════════════════════════════════════════════
h1(doc, "Fonctionnalité 1 : Accès et Navigation")

# F1.1 ──────────────────────────────────────────────────────────────────
h2(doc, "F1.1 : Affichage du menu « Chéquier »")
h3(doc, "1. Définition et Objectif")
body(doc, "Cette fonctionnalité permet au client authentifié sur la plateforme Internet Banking d'accéder à l'espace dédié à la gestion de ses demandes de chéquier. Un nouveau menu « Chéquier » est ajouté à la barre de navigation principale, organisé en deux sous-menus distincts : « Mes Demandes » pour consulter l'historique et le suivi des demandes existantes, et « Nouvelle Demande » pour initier une nouvelle demande. L'objectif est d'offrir un point d'entrée unique, clair et intuitif, accessible depuis tout appareil (web ou mobile), sans nécessiter de déplacement en agence.")
h3(doc, "2. Variantes")
bul(doc, "Accès depuis un navigateur web (desktop)")
bul(doc, "Accès depuis l'application mobile Internet Banking")
bul(doc, "Client sans aucun compte éligible (affichage du menu mais accès restreint)")
h3(doc, "3. Entrées")
bul(doc, "Session client active et authentifiée (login + mot de passe validés)")
h3(doc, "4. Sorties")
bul(doc, "Affichage du menu « Chéquier » dans la barre de navigation principale")
bul(doc, "Sous-menu « Mes Demandes » : redirige vers la liste des demandes du client")
bul(doc, "Sous-menu « Nouvelle Demande » : redirige vers le formulaire de soumission")
h3(doc, "5. Préconditions")
bul(doc, "Client authentifié avec une session valide et non expirée sur la plateforme Internet Banking")
h3(doc, "6. Postconditions")
bul(doc, "Le client visualise les deux sous-menus disponibles")
bul(doc, "La navigation est journalisée (identifiant client, horodatage, action)")
h3(doc, "7. Processus")
num(doc, "Le client se connecte à la plateforme Internet Banking AFBSS avec ses identifiants.")
num(doc, "Le système valide la session et affiche le tableau de bord principal.")
num(doc, "Le client repère et clique sur le menu « Chéquier » dans la barre de navigation.")
num(doc, "Le système affiche les deux sous-menus : « Mes Demandes » et « Nouvelle Demande ».")
num(doc, "Le client sélectionne l'action souhaitée.")
num(doc, "Le système journalise l'accès avec l'identifiant client et l'horodatage.")
h3(doc, "8. Scénarios Alternatifs et Cas d'Erreur")
bul(doc, "Session expirée → Le système redirige automatiquement vers la page de connexion avec le message : « Votre session a expiré. Veuillez vous reconnecter pour continuer. »")
bul(doc, "Client sans compte éligible → Le sous-menu « Nouvelle Demande » est visible mais affiche à l'ouverture : « Aucun de vos comptes n'est éligible à une demande de chéquier en ligne. Veuillez contacter votre agence. »")
empty(doc)

# F1.2 ──────────────────────────────────────────────────────────────────
h2(doc, "F1.2 : Consultation de l'historique des demandes (« Mes Demandes »)")
h3(doc, "1. Définition et Objectif")
body(doc, "Cette fonctionnalité permet au client de consulter l'ensemble de ses demandes de chéquier passées et en cours, avec leur statut actualisé en temps réel. Elle offre une visibilité complète et transparente sur l'avancement de chaque demande, sans avoir à contacter l'agence. Le client peut également consulter le détail d'une demande spécifique (compte débité, frais appliqués, agence de livraison, historique des transitions de statut).")
h3(doc, "2. Variantes")
bul(doc, "Client ayant plusieurs demandes en cours simultanément")
bul(doc, "Client consultant une demande clôturée (statut « Livré »)")
bul(doc, "Client consultant une demande échouée (statut « Échoué »)")
h3(doc, "3. Entrées")
bul(doc, "Identifiant client (extrait de la session active)")
h3(doc, "4. Sorties")
body(doc, "Liste des demandes triées par date de soumission décroissante, affichant pour chaque demande :")
bul(doc, "Numéro de référence unique", 1)
bul(doc, "Date et heure de soumission", 1)
bul(doc, "Numéro de compte débité", 1)
bul(doc, "Devise et nombre de feuillets demandés", 1)
bul(doc, "Agence de livraison sélectionnée", 1)
bul(doc, "Statut actuel (avec badge coloré) et date du dernier changement de statut", 1)
body(doc, "Vue détaillée accessible par clic sur une demande, incluant :")
bul(doc, "Détail des frais appliqués (commission + taxes)", 1)
bul(doc, "Historique complet des transitions de statut avec horodatage", 1)
h3(doc, "5. Préconditions")
bul(doc, "Client authentifié avec une session valide")
bul(doc, "Base de données Internet Banking accessible")
h3(doc, "6. Postconditions")
bul(doc, "Le client visualise ses demandes avec les statuts actualisés en temps réel")
bul(doc, "Aucune modification des données n'est possible depuis cet écran")
h3(doc, "7. Processus")
num(doc, "Le client sélectionne « Mes Demandes » dans le menu « Chéquier ».")
num(doc, "Le système interroge la base de données et récupère toutes les demandes associées à l'identifiant client.")
num(doc, "La liste est affichée avec les statuts actualisés, triée par date décroissante.")
num(doc, "Le client peut cliquer sur une demande pour en afficher le détail complet.")
num(doc, "Le système affiche le détail : informations de la demande, frais, historique des transitions.")
h3(doc, "8. Scénarios Alternatifs et Cas d'Erreur")
bul(doc, "Aucune demande existante → Affichage du message : « Vous n'avez pas encore effectué de demande de chéquier en ligne. Cliquez sur \"Nouvelle Demande\" pour commencer. »")
bul(doc, "Indisponibilité temporaire → Message : « Impossible de charger vos demandes pour le moment. Veuillez réessayer dans quelques instants. »")

# ════════════════════════════════════════════════════════════════════════
# FONCTIONNALITÉ 2
# ════════════════════════════════════════════════════════════════════════
h1(doc, "Fonctionnalité 2 : Soumission de la Demande")

# F2.1 ──────────────────────────────────────────────────────────────────
h2(doc, "F2.1 : Saisie du formulaire de demande")
h3(doc, "1. Définition et Objectif")
body(doc, "Cette fonctionnalité permet au client de renseigner les informations nécessaires à sa demande de chéquier via un formulaire structuré, guidé et sécurisé. Elle garantit que seuls les comptes éligibles sont proposés à la sélection, que la devise est automatiquement renseignée selon le compte choisi, et que toutes les données obligatoires sont présentes et cohérentes avant de procéder à l'étape de confirmation. Le formulaire est conçu pour être complété en moins de 2 minutes, avec un minimum de saisie manuelle.")
h3(doc, "2. Variantes")
bul(doc, "Client possédant un seul compte éligible (pré-sélection automatique du compte)")
bul(doc, "Client possédant plusieurs comptes éligibles (sélection manuelle dans la liste déroulante)")
bul(doc, "Client choisissant une agence de livraison différente de son agence de domiciliation")
bul(doc, "Client choisissant son agence de domiciliation comme agence de livraison")
h3(doc, "3. Entrées")
bul(doc, "Sélection du compte à débiter (liste déroulante des comptes éligibles du client, récupérée via le Core Banking)")
bul(doc, "Devise (champ auto-renseigné et non modifiable, déterminée par le compte sélectionné)")
bul(doc, "Nombre de feuillets : 25 ou 50 (sélection via boutons radio, valeur par défaut : 25)")
bul(doc, "Agence de livraison (liste déroulante de toutes les agences AFBSS actives, issue du référentiel agences)")
h3(doc, "4. Sorties")
bul(doc, "Formulaire validé avec toutes les données obligatoires renseignées")
bul(doc, "Données temporairement stockées en session (aucune écriture en base à ce stade)")
bul(doc, "Calcul des frais déclenché automatiquement (F2.2)")
bul(doc, "Redirection vers l'écran de confirmation (F2.3)")
h3(doc, "5. Préconditions")
bul(doc, "Client authentifié avec au moins un compte actif, non bloqué et disposant de l'autorisation chéquier dans le Core Banking (BR-01)")
bul(doc, "Référentiel des agences AFBSS disponible et à jour")
bul(doc, "Barème de frais disponible dans le Core Banking")
h3(doc, "6. Postconditions")
bul(doc, "Les données du formulaire sont stockées en session de manière sécurisée")
bul(doc, "Aucune opération comptable ni écriture en base de données à ce stade")
bul(doc, "Un jeton de session unique (token d'idempotence) est généré pour cette instance de formulaire (BR-10)")
h3(doc, "7. Processus")
num(doc, "Le client accède au formulaire via le sous-menu « Nouvelle Demande ».")
num(doc, "Le système appelle le Core Banking pour récupérer la liste des comptes éligibles du client.")
num(doc, "Si un seul compte éligible : il est pré-sélectionné automatiquement. Si plusieurs : liste déroulante affichée.")
num(doc, "Dès la sélection du compte, le système auto-renseigne la devise correspondante (non modifiable).")
num(doc, "Le client choisit le nombre de feuillets (25 ou 50).")
num(doc, "Le client sélectionne l'agence de livraison souhaitée dans la liste déroulante.")
num(doc, "Le système calcule automatiquement les frais en arrière-plan (F2.2).")
num(doc, "Le client clique sur « Continuer ».")
num(doc, "Le système vérifie que tous les champs obligatoires sont renseignés.")
num(doc, "Si validation OK : redirection vers l'écran de confirmation (F2.3).")
num(doc, "Le système génère et associe un jeton d'idempotence unique à cette session de formulaire.")
h3(doc, "8. Scénarios Alternatifs et Cas d'Erreur")
bul(doc, "Aucun compte éligible → Le formulaire n'est pas affiché. Message : « Aucun de vos comptes n'est éligible à une demande de chéquier en ligne. »")
bul(doc, "Champ obligatoire manquant → Mise en évidence visuelle du champ avec un message d'erreur inline explicite.")
bul(doc, "Indisponibilité du référentiel agences → Message : « Impossible de charger la liste des agences. Veuillez réessayer. »")
bul(doc, "Indisponibilité du Core Banking → Message : « Service temporairement indisponible. Veuillez réessayer. »")
empty(doc)

# F2.2 ──────────────────────────────────────────────────────────────────
h2(doc, "F2.2 : Calcul et affichage automatique des frais")
h3(doc, "1. Définition et Objectif")
body(doc, "Cette sous-fonctionnalité assure le calcul automatique, transparent et en temps réel des frais applicables à la demande de chéquier (commission bancaire + taxes réglementaires), sur la base du barème configuré dans le Core Banking ALTBank. Les frais sont calculés dès que le compte et le nombre de feuillets sont sélectionnés dans le formulaire, et affichés de manière détaillée sur l'écran de confirmation avant toute validation.")
h3(doc, "2. Variantes")
bul(doc, "Calcul pour un chéquier de 25 feuillets")
bul(doc, "Calcul pour un chéquier de 50 feuillets")
bul(doc, "Barème différencié selon la devise du compte (SSP, USD, etc.)")
h3(doc, "3. Entrées")
bul(doc, "Identifiant du compte sélectionné (pour déterminer la devise et l'agence de domiciliation)")
bul(doc, "Nombre de feuillets choisi (25 ou 50)")
h3(doc, "4. Sorties")
bul(doc, "Montant de la commission bancaire (en devise du compte)")
bul(doc, "Montant des taxes applicables (en devise du compte)")
bul(doc, "Montant total des frais = Commission + Taxes (affiché sur l'écran de confirmation)")
h3(doc, "5. Préconditions")
bul(doc, "Barème de frais disponible, actif et à jour dans le Core Banking ALTBank")
bul(doc, "Compte sélectionné valide avec devise identifiée")
h3(doc, "6. Postconditions")
bul(doc, "Les frais calculés sont transmis et affichés sur l'écran de confirmation (F2.3)")
bul(doc, "Le montant total est conservé en session pour être utilisé lors de l'exécution de la transaction (F4.1)")
h3(doc, "7. Processus")
num(doc, "Dès que le client sélectionne un compte et un nombre de feuillets, le système interroge le Core Banking pour récupérer le barème applicable.")
num(doc, "Le système calcule : Frais totaux = Commission + Taxes.")
num(doc, "Les trois montants (commission, taxes, total) sont conservés en session.")
num(doc, "Ils sont affichés de manière détaillée et lisible sur l'écran de confirmation (F2.3).")
h3(doc, "8. Scénarios Alternatifs et Cas d'Erreur")
bul(doc, "Barème non disponible → Blocage de la progression avec le message : « Impossible de calculer les frais pour le moment. Veuillez réessayer ultérieurement. »")
bul(doc, "Barème renvoyé incohérent → Alerte technique journalisée + blocage + message client : « Une erreur est survenue lors du calcul des frais. Veuillez contacter votre agence. »")
empty(doc)

# F2.3 ──────────────────────────────────────────────────────────────────
h2(doc, "F2.3 : Écran de confirmation récapitulatif")
h3(doc, "1. Définition et Objectif")
body(doc, "Avant toute action irréversible (envoi OTP, débit du compte), cet écran présente au client un récapitulatif complet et lisible de sa demande, incluant toutes les informations saisies et les frais calculés. Il constitue la dernière opportunité pour le client de vérifier, corriger ou annuler sa demande avant engagement, conformément aux exigences de consentement éclairé (BR-06).")
h3(doc, "2. Variantes")
bul(doc, "Client confirmant sa demande → déclenchement de l'OTP (F3.1)")
bul(doc, "Client souhaitant modifier ses informations → retour au formulaire (F2.1)")
bul(doc, "Client annulant sa demande → retour au tableau de bord")
h3(doc, "3. Entrées")
bul(doc, "Données saisies dans le formulaire (F2.1) : compte, devise, nombre de feuillets, agence de livraison")
bul(doc, "Frais calculés (F2.2) : commission, taxes, total")
h3(doc, "4. Sorties")
body(doc, "Écran récapitulatif affichant :")
bul(doc, "Numéro et libellé du compte à débiter", 1)
bul(doc, "Solde disponible actuel du compte", 1)
bul(doc, "Devise et nombre de feuillets demandé", 1)
bul(doc, "Agence de livraison sélectionnée (nom et adresse)", 1)
bul(doc, "Détail des frais : commission, taxes, total à débiter", 1)
bul(doc, "Mention explicite : « Ce montant sera immédiatement débité de votre compte lors de la validation. »", 1)
bul(doc, "Bouton « Confirmer et recevoir mon code OTP »", 1)
bul(doc, "Bouton « Modifier » (retour au formulaire)", 1)
bul(doc, "Bouton « Annuler » (retour au tableau de bord, sans aucune action)", 1)
h3(doc, "5. Préconditions")
bul(doc, "Formulaire correctement renseigné et validé (F2.1)")
bul(doc, "Frais calculés avec succès (F2.2)")
h3(doc, "6. Postconditions")
bul(doc, "Si le client clique sur « Confirmer » : déclenchement de l'envoi OTP (F3.1)")
bul(doc, "Si le client clique sur « Modifier » : retour au formulaire avec les données pré-remplies")
bul(doc, "Si le client clique sur « Annuler » : invalidation du jeton de session, retour au tableau de bord")
h3(doc, "7. Processus")
num(doc, "Le système affiche l'écran récapitulatif avec toutes les informations de la demande et les frais calculés.")
num(doc, "Le client prend connaissance du résumé et des frais qui seront débités.")
num(doc, "Le client clique sur « Confirmer et recevoir mon code OTP » pour procéder à la validation.")
num(doc, "Le système déclenche la génération et l'envoi de l'OTP (F3.1).")

# ════════════════════════════════════════════════════════════════════════
# FONCTIONNALITÉ 3
# ════════════════════════════════════════════════════════════════════════
h1(doc, "Fonctionnalité 3 : Validation et Sécurisation")

# F3.1 ──────────────────────────────────────────────────────────────────
h2(doc, "F3.1 : Génération et envoi de l'OTP")
h3(doc, "1. Définition et Objectif")
body(doc, "Cette fonctionnalité assure la sécurisation transactionnelle de la demande de chéquier en générant et en envoyant un code OTP (One-Time Password) à usage unique au client, via son canal de communication préféré (SMS en priorité, email en secondaire). L'OTP constitue le second facteur d'authentification transactionnelle obligatoire pour toute soumission de demande, conformément aux exigences de sécurité bancaire (BR-06).")
h3(doc, "2. Variantes")
bul(doc, "Envoi par SMS (canal prioritaire)")
bul(doc, "Envoi par email (canal secondaire si SMS indisponible ou selon préférence client)")
bul(doc, "Renvoi d'un nouvel OTP à la demande du client (en cas d'expiration)")
h3(doc, "3. Entrées")
bul(doc, "Confirmation de l'écran récapitulatif par le client (clic sur « Confirmer »)")
bul(doc, "Numéro de téléphone mobile et/ou adresse email du client (issus du profil client)")
bul(doc, "Canal de communication préféré du client")
h3(doc, "4. Sorties")
bul(doc, "OTP généré : code numérique à 6 chiffres, usage unique, validité de 5 minutes")
bul(doc, "Message envoyé au client contenant : le code OTP, le rappel du montant à débiter, la durée de validité")
bul(doc, "Écran de saisie OTP affiché avec un compteur de temps visible (décompte en temps réel)")
bul(doc, "Lien « Renvoyer un code » disponible après expiration")
h3(doc, "5. Préconditions")
bul(doc, "Récapitulatif confirmé par le client")
bul(doc, "Coordonnées de contact valides et disponibles dans le profil client")
bul(doc, "Passerelle SMS et/ou service email opérationnels")
h3(doc, "6. Postconditions")
bul(doc, "L'OTP est généré, hashé et stocké temporairement côté serveur avec son horodatage de création")
bul(doc, "Le compteur de tentatives est initialisé à 0")
bul(doc, "Le chronomètre de validité (5 minutes) est démarré")
bul(doc, "L'événement est journalisé (génération OTP, canal utilisé, horodatage)")
h3(doc, "7. Processus")
num(doc, "Le client clique sur « Confirmer et recevoir mon code OTP » sur l'écran récapitulatif.")
num(doc, "Le système génère un code OTP aléatoire à 6 chiffres.")
num(doc, "Le code est hashé et stocké temporairement en session serveur avec l'horodatage de génération.")
num(doc, "Le système envoie le message via le canal prioritaire du client (SMS).")
num(doc, "En cas d'échec d'envoi SMS, tentative automatique par email, avec journalisation de l'incident.")
num(doc, "L'écran de saisie OTP est affiché avec un compteur de décompte visible.")
num(doc, "L'événement (génération, canal, horodatage) est journalisé dans le journal d'audit.")
h3(doc, "8. Scénarios Alternatifs et Cas d'Erreur")
bul(doc, "Échec d'envoi SMS → Tentative automatique par email + journalisation + message au client : « Votre code vous a été envoyé par email. »")
bul(doc, "Coordonnées absentes ou invalides → Message : « Impossible d'envoyer le code de validation. Vos coordonnées sont manquantes. Veuillez contacter votre agence. »")
bul(doc, "OTP expiré → Affichage du lien « Renvoyer un code » + message : « Votre code a expiré. Cliquez sur \"Renvoyer un code\" pour en recevoir un nouveau. »")
empty(doc)

# F3.2 ──────────────────────────────────────────────────────────────────
h2(doc, "F3.2 : Vérification de l'OTP et autorisation de la transaction")
h3(doc, "1. Définition et Objectif")
body(doc, "Cette fonctionnalité vérifie la validité du code OTP saisi par le client, contrôle son intégrité (concordance, non-expiration, nombre de tentatives) et autorise ou bloque le passage à l'étape d'exécution Core Banking. Elle intègre un mécanisme de protection contre les attaques par force brute via un compteur de tentatives limité à 3.")
h3(doc, "2. Variantes")
bul(doc, "OTP correct à la première tentative")
bul(doc, "OTP incorrect avec tentatives restantes")
bul(doc, "OTP expiré (délai de 5 minutes dépassé)")
bul(doc, "Tentatives épuisées (3 échecs consécutifs)")
h3(doc, "3. Entrées")
bul(doc, "Code OTP saisi par le client (6 chiffres)")
bul(doc, "Référence de la session de demande en cours (jeton de session)")
bul(doc, "Horodatage de génération de l'OTP (pour contrôle d'expiration)")
bul(doc, "Compteur de tentatives actuel")
h3(doc, "4. Sorties")
bul(doc, "OTP valide et non expiré : déclenchement de F4.1 (appel Core Banking et exécution atomique)")
bul(doc, "OTP invalide avec tentatives restantes : message d'erreur + décrémentation du compteur")
bul(doc, "OTP expiré : annulation de la saisie + proposition de renvoi")
bul(doc, "Tentatives épuisées (3 échecs) : invalidation définitive de la session de demande")
h3(doc, "5. Préconditions")
bul(doc, "OTP généré et en cours de validité (moins de 5 minutes depuis la génération)")
bul(doc, "Session de demande active et non invalidée")
h3(doc, "6. Postconditions")
bul(doc, "Si OTP valide : processus d'exécution Core Banking déclenché (F4.1)")
bul(doc, "Si échec définitif : session invalidée, jeton d'idempotence détruit, client invité à recommencer")
h3(doc, "7. Règles de Gestion")
bul(doc, "Maximum 3 tentatives de saisie par OTP généré")
bul(doc, "Expiration automatique de l'OTP après 5 minutes")
bul(doc, "Après 3 échecs : blocage définitif avec le message : « Trop de tentatives incorrectes. Votre demande a été annulée pour des raisons de sécurité. »")
bul(doc, "Tous les échecs de saisie sont journalisés avec identifiant client, horodatage et nombre de tentatives")
h3(doc, "8. Processus")
num(doc, "Le client saisit le code OTP reçu sur l'écran de validation.")
num(doc, "Le système compare le hash du code saisi avec le hash stocké en session.")
num(doc, "Le système vérifie que le délai de validité (5 minutes) n'est pas dépassé.")
num(doc, "Si le code est correct et non expiré : le système valide l'OTP et déclenche F4.1.")
num(doc, "Si le code est incorrect : le compteur de tentatives est incrémenté. Si tentatives < 3 : message d'erreur. Si tentatives = 3 : invalidation de la session.")
num(doc, "Toute tentative (réussie ou non) est journalisée.")
h3(doc, "9. Scénarios Alternatifs et Cas d'Erreur")
bul(doc, "OTP expiré → Message : « Votre code de validation a expiré. » + Bouton « Recevoir un nouveau code ».")
bul(doc, "OTP incorrect (1re ou 2e tentative) → Message : « Code incorrect. Il vous reste X tentative(s). »")
bul(doc, "3 échecs consécutifs → Message : « Trop de tentatives incorrectes. Votre demande a été annulée. »")

# ════════════════════════════════════════════════════════════════════════
# FONCTIONNALITÉ 4
# ════════════════════════════════════════════════════════════════════════
h1(doc, "Fonctionnalité 4 : Exécution et Création de la Demande")

# F4.1 ──────────────────────────────────────────────────────────────────
h2(doc, "F4.1 : Appel au Core Banking et exécution de la transaction atomique")
h3(doc, "1. Définition et Objectif")
body(doc, "C'est la fonctionnalité centrale et critique du module M-CHQ. Suite à la validation réussie de l'OTP, l'Internet Banking déclenche une transaction atomique dans le Core Banking ALTBank, comprenant simultanément et indivisiblement : le débit du compte client du montant total des frais (commission + taxes), et le crédit des comptes de commission et de taxes de l'agence de domiciliation du compte (Home Branch). Si l'une des deux opérations ne peut pas être exécutée, l'ensemble de la transaction est annulé sans aucun effet partiel (principe du tout ou rien - BR-02).")
note_critical(doc, "Règle critique (BR-04)", "Les comptes crédités (commission et taxes) sont obligatoirement ceux de l'agence de domiciliation du compte débité (Home Branch), quelle que soit l'agence de livraison choisie par le client. Cette règle est non négociable et doit être strictement contrôlée en tests de recette.")
h3(doc, "2. Variantes")
bul(doc, "Transaction exécutée avec succès (solde suffisant, Core Banking disponible)")
bul(doc, "Transaction refusée pour solde insuffisant")
bul(doc, "Transaction échouée pour indisponibilité ou timeout du Core Banking")
bul(doc, "Transaction refusée pour compte bloqué ou non éligible")
h3(doc, "3. Entrées")
bul(doc, "Identifiant du compte client à débiter")
bul(doc, "Montant total des frais calculés et validés (commission + taxes)")
bul(doc, "Identifiant de l'agence de domiciliation du compte (Home Branch) — pour le crédit comptable")
bul(doc, "Jeton d'idempotence de session (pour éviter la double exécution)")
bul(doc, "Référence de la demande (générée en session)")
h3(doc, "4. Sorties")
bul(doc, "Succès : accusé de transaction du Core Banking avec le numéro de référence, la date et l'heure d'exécution, la confirmation du débit et du crédit effectués")
bul(doc, "Échec solde insuffisant : code erreur INSUFFICIENT_FUNDS + motif explicite")
bul(doc, "Échec technique : code erreur CORE_BANKING_UNAVAILABLE ou TIMEOUT + motif")
bul(doc, "Échec compte : code erreur ACCOUNT_BLOCKED ou ACCOUNT_NOT_ELIGIBLE + motif")
h3(doc, "5. Préconditions")
bul(doc, "OTP validé avec succès (F3.2)")
bul(doc, "Connexion au Core Banking ALTBank disponible et opérationnelle")
bul(doc, "Jeton d'idempotence valide et non déjà utilisé")
h3(doc, "6. Postconditions")
bul(doc, "En cas de succès : débit effectif du compte client + crédit effectif des comptes de commission et taxes de la Home Branch + enregistrement dans la table checkbook_requests")
bul(doc, "En cas d'échec : aucun débit, aucun crédit, aucun enregistrement partiel dans le Core Banking")
bul(doc, "Dans tous les cas : journalisation complète de la tentative (paramètres, réponse, horodatage)")
h3(doc, "7. Processus")
num(doc, "L'Internet Banking envoie une requête de transaction sécurisée au Core Banking ALTBank, contenant : l'identifiant du compte à débiter, le montant total des frais, l'identifiant de la Home Branch et le jeton d'idempotence.")
num(doc, "Le Core Banking vérifie l'idempotence (la transaction n'a pas déjà été exécutée avec ce jeton).")
num(doc, "Le Core Banking vérifie que le compte est actif, non bloqué et dispose d'un solde suffisant.")
num(doc, "Si toutes les vérifications sont positives : exécution atomique (débit + crédit simultanés) en une seule opération indivisible.")
num(doc, "Le Core Banking enregistre la demande dans sa table dédiée (checkbook_requests).")
num(doc, "Le Core Banking retourne un accusé de succès à l'Internet Banking avec le numéro de référence de la transaction.")
num(doc, "L'Internet Banking reçoit l'accusé de succès et déclenche immédiatement F4.2 (création du ticket IB).")
h3(doc, "8. Scénarios Alternatifs et Cas d'Erreur")
bul(doc, "Solde insuffisant → Message : « Votre solde est insuffisant pour couvrir les frais. Aucun montant n'a été débité. Veuillez alimenter votre compte et réessayer. »")
bul(doc, "Core Banking indisponible ou timeout → Message : « Le service bancaire est temporairement indisponible. Votre demande n'a pas été enregistrée et aucun montant n'a été débité. »")
bul(doc, "Compte bloqué ou non éligible → Message : « Votre compte ne permet pas d'effectuer cette opération. Veuillez contacter votre agence. »")
bul(doc, "Jeton d'idempotence déjà utilisé (double soumission) → Rejet silencieux côté serveur + retour au statut de la demande déjà créée.")
empty(doc)

# F4.2 ──────────────────────────────────────────────────────────────────
h2(doc, "F4.2 : Création du ticket Internet Banking")
h3(doc, "1. Définition et Objectif")
body(doc, "Uniquement et exclusivement après réception d'un accusé de succès formel du Core Banking ALTBank (BR-08), l'Internet Banking procède à la création de son propre enregistrement de la demande (ticket IB). Cette séquence stricte — Core Banking en premier, ticket IB en second — garantit l'invariant fondamental du module : il ne peut jamais exister une demande enregistrée dans l'Internet Banking sans transaction comptable correspondante et confirmée dans ALTBank.")
h3(doc, "2. Entrées")
bul(doc, "Accusé de succès du Core Banking (numéro de référence de la transaction, horodatage d'exécution)")
bul(doc, "Données complètes de la demande (compte, devise, nombre de feuillets, agence de livraison)")
bul(doc, "Identifiant client")
bul(doc, "Montant des frais débités (commission + taxes)")
h3(doc, "3. Sorties")
body(doc, "Ticket de demande créé dans la base de données Internet Banking, contenant :")
bul(doc, "Numéro de référence IB unique (ex. : CHQ-2026-XXXXXX)", 1)
bul(doc, "Référence de la transaction Core Banking (lien avec ALTBank)", 1)
bul(doc, "Statut initial : « En attente de traitement »", 1)
bul(doc, "Date et heure de création", 1)
bul(doc, "Identifiant client et numéro de compte débité", 1)
bul(doc, "Devise et nombre de feuillets", 1)
bul(doc, "Agence de livraison sélectionnée", 1)
bul(doc, "Montant total des frais débités", 1)
bul(doc, "Notification de succès envoyée au client (F6.1)")
bul(doc, "Demande visible dans « Mes Demandes » côté client")
bul(doc, "Demande visible dans la file d'attente de l'agence de livraison côté back-office opérateur")
h3(doc, "4. Préconditions")
bul(doc, "Accusé de succès du Core Banking reçu et validé")
bul(doc, "Base de données Internet Banking disponible")
h3(doc, "5. Postconditions")
bul(doc, "La demande est enregistrée en base IB avec le statut « En attente de traitement »")
bul(doc, "La demande apparaît dans l'interface client (« Mes Demandes »)")
bul(doc, "La demande apparaît dans la file d'attente de l'agence de livraison (back-office)")
bul(doc, "La création est journalisée (numéro de référence, horodatage, identifiant client)")
h3(doc, "6. Processus")
num(doc, "L'Internet Banking reçoit l'accusé de succès du Core Banking avec le numéro de référence de transaction.")
num(doc, "Le système génère un numéro de référence IB unique pour la demande.")
num(doc, "Le système crée l'enregistrement complet de la demande en base de données IB.")
num(doc, "Le statut initial « En attente de traitement » est affecté automatiquement.")
num(doc, "La création est journalisée.")
num(doc, "Le système déclenche l'envoi de la notification de succès au client (F6.1).")
num(doc, "La demande devient immédiatement visible dans l'interface client et dans le back-office de l'agence de livraison.")
h3(doc, "7. Scénarios Alternatifs et Cas d'Erreur")
bul(doc, "Échec d'écriture en base IB après succès Core Banking → Incident critique journalisé immédiatement avec alerte au niveau administrateur. Un mécanisme de réconciliation doit être déclenché pour créer le ticket IB en rattrapage.")
empty(doc)

# F4.3 ──────────────────────────────────────────────────────────────────
h2(doc, "F4.3 : Gestion des cas d'échec")
h3(doc, "1. Définition et Objectif")
body(doc, "Cette fonctionnalité définit le comportement du système en cas d'échec à toute étape du processus de soumission. Elle garantit qu'en toutes circonstances, aucune situation intermédiaire incohérente n'est créée dans le système (pas de débit sans ticket IB, pas de ticket IB sans débit confirmé - BR-02 et BR-08), et que le client reçoit un message clair, explicite et orienté action.")
empty(doc)
make_table(doc,
    ["Cas d'erreur", "Comportement système", "Message affiché au client", "Action proposée"],
    [
        ["Solde insuffisant", "Aucun débit, aucune demande créée", "Votre solde est insuffisant. Aucun montant n'a été débité.", "Alimentez votre compte et réessayez."],
        ["Core Banking indisponible / timeout", "Annulation complète, aucune opération partielle", "Service temporairement indisponible. Votre demande n'a pas été enregistrée.", "Réessayez dans quelques instants."],
        ["OTP invalide (tentatives restantes)", "Maintien sur l'écran OTP, compteur décrémenté", "Code incorrect. Il vous reste X tentative(s).", "Ressaisie du code OTP"],
        ["OTP expiré", "Invalidation du code, maintien de la session", "Votre code a expiré.", "Cliquez sur « Renvoyer un code »."],
        ["3 échecs OTP consécutifs", "Invalidation définitive de la session de demande", "Trop de tentatives. Votre demande a été annulée.", "Recommencez votre demande."],
        ["Double soumission", "Rejet par contrôle d'idempotence côté serveur", "Cette demande a déjà été soumise.", "Redirection vers « Mes Demandes »"],
        ["Compte bloqué / inéligible", "Blocage dès le formulaire", "Votre compte ne permet pas cette opération.", "Contactez votre agence."],
    ],
    [3.8, 4.0, 4.8, 3.6]
)

# ════════════════════════════════════════════════════════════════════════
# FONCTIONNALITÉ 5
# ════════════════════════════════════════════════════════════════════════
h1(doc, "Fonctionnalité 5 : Traitement Back-Office")

# F5.1 ──────────────────────────────────────────────────────────────────
h2(doc, "F5.1 : Consultation de la file d'attente (opérateur)")
h3(doc, "1. Définition et Objectif")
body(doc, "L'interface back-office dédiée aux opérateurs d'agence affiche exclusivement les demandes de chéquier dont l'agence de livraison correspond à l'agence d'appartenance de l'opérateur connecté. Ce cloisonnement strict (BR-09) garantit que chaque opérateur ne traite que les demandes relevant de sa responsabilité. L'interface offre des fonctionnalités de filtrage et de tri pour faciliter la gestion quotidienne.")
h3(doc, "2. Variantes")
bul(doc, "File d'attente avec plusieurs demandes en statut « En attente de traitement »")
bul(doc, "File d'attente vide (aucune demande en attente)")
bul(doc, "Consultation des demandes dans tous les statuts (historique complet)")
h3(doc, "3. Entrées")
bul(doc, "Authentification de l'opérateur (identifiant et agence d'appartenance extraits du profil)")
h3(doc, "4. Sorties")
body(doc, "Liste des demandes filtrées par agence de livraison de l'opérateur, affichant pour chaque demande :")
bul(doc, "Numéro de référence IB", 1)
bul(doc, "Date et heure de soumission", 1)
bul(doc, "Nom du client et numéro de compte", 1)
bul(doc, "Devise et nombre de feuillets", 1)
bul(doc, "Statut actuel avec horodatage du dernier changement", 1)
bul(doc, "Opérateur responsable (si déjà pris en charge)", 1)
bul(doc, "Filtres disponibles : par statut, par date, par numéro de compte", 1)
bul(doc, "Tri par défaut : date de soumission croissante (plus ancienne demande en premier)", 1)
h3(doc, "5. Préconditions")
bul(doc, "Opérateur authentifié avec le profil « Opérateur agence »")
bul(doc, "Agence d'appartenance de l'opérateur identifiée dans son profil")
h3(doc, "6. Postconditions")
bul(doc, "L'opérateur visualise uniquement les demandes de son agence de livraison")
bul(doc, "Aucune modification des données n'est effectuée à cette étape")
empty(doc)

# F5.2 ──────────────────────────────────────────────────────────────────
h2(doc, "F5.2 : Prise en charge d'une demande (→ statut « En cours »)")
h3(doc, "1. Définition et Objectif")
body(doc, "Cette action permet à un opérateur d'agence de prendre officiellement en charge une demande dont le statut est « En attente de traitement », signalant ainsi au système, aux autres opérateurs et au client que le traitement physique du chéquier a débuté.")
h3(doc, "2. Entrées")
bul(doc, "Sélection de la demande par l'opérateur dans la file d'attente")
bul(doc, "Action « Prendre en charge » (bouton d'action sur la demande sélectionnée)")
h3(doc, "3. Sorties")
bul(doc, "Statut de la demande mis à jour : « En attente de traitement » → « En cours »")
bul(doc, "Identifiant de l'opérateur et horodatage de la prise en charge enregistrés")
bul(doc, "Notification automatique envoyée au client (F6.3) : « Votre demande de chéquier est en cours de traitement par notre équipe. »")
bul(doc, "Journalisation de l'action (opérateur, demande, horodatage)")
h3(doc, "4. Préconditions")
bul(doc, "Demande en statut « En attente de traitement »")
bul(doc, "Opérateur appartenant à l'agence de livraison de la demande (contrôle RBAC)")
h3(doc, "5. Postconditions")
bul(doc, "Le statut « En cours » est visible côté client dans « Mes Demandes »")
bul(doc, "La demande est marquée comme prise en charge par l'opérateur dans le back-office")
h3(doc, "6. Scénarios Alternatifs et Cas d'Erreur")
bul(doc, "Tentative de prise en charge d'une demande déjà prise par un autre opérateur → Message : « Cette demande est déjà prise en charge par [Opérateur X]. »")
empty(doc)

# F5.3 ──────────────────────────────────────────────────────────────────
h2(doc, "F5.3 : Déclaration de l'impression (→ statut « Imprimé / Prêt »)")
h3(doc, "1. Définition et Objectif")
body(doc, "L'opérateur déclare dans le système que le chéquier a été physiquement imprimé et est disponible pour retrait par le client dans l'agence de livraison. Cette action déclenche la notification la plus importante pour le client.")
h3(doc, "2. Entrées")
bul(doc, "Action « Déclarer Imprimé / Prêt » sur la demande en statut « En cours »")
h3(doc, "3. Sorties")
bul(doc, "Statut mis à jour : « En cours » → « Imprimé / Prêt »")
bul(doc, "Identifiant de l'opérateur et horodatage enregistrés")
bul(doc, "Notification automatique envoyée au client : « Votre chéquier est prêt. Présentez-vous à l'agence [Nom et adresse] muni de votre pièce d'identité pour le retirer. »")
bul(doc, "Journalisation de l'action")
h3(doc, "4. Préconditions")
bul(doc, "Demande en statut « En cours »")
bul(doc, "Opérateur responsable de la demande ou appartenant à l'agence de livraison")
empty(doc)

# F5.4 ──────────────────────────────────────────────────────────────────
h2(doc, "F5.4 : Confirmation de livraison (→ statut « Livré »)")
h3(doc, "1. Définition et Objectif")
body(doc, "L'opérateur confirme dans le système la remise physique du chéquier au client après vérification de son identité au guichet. Cette action clôture définitivement le cycle de vie de la demande et déclenche la notification finale au client.")
h3(doc, "2. Entrées")
bul(doc, "Action « Confirmer Livré » sur la demande en statut « Imprimé / Prêt »")
h3(doc, "3. Sorties")
bul(doc, "Statut mis à jour : « Imprimé / Prêt » → « Livré »")
bul(doc, "Identifiant de l'opérateur et horodatage de livraison enregistrés")
bul(doc, "Notification finale envoyée au client : « Votre chéquier vous a été remis. Merci pour votre confiance. AFBSS reste à votre disposition. »")
bul(doc, "Demande définitivement clôturée dans le système (état terminal)")
bul(doc, "Journalisation complète de la clôture")
h3(doc, "4. Préconditions")
bul(doc, "Demande en statut « Imprimé / Prêt »")
bul(doc, "Opérateur appartenant à l'agence de livraison de la demande")

# ════════════════════════════════════════════════════════════════════════
# FONCTIONNALITÉ 6
# ════════════════════════════════════════════════════════════════════════
h1(doc, "Fonctionnalité 6 : Notifications Automatiques")
h2(doc, "F6.1 – F6.3 : Notifications client à chaque étape")
h3(doc, "1. Définition et Objectif")
body(doc, "Le système envoie automatiquement des notifications au client à chaque étape significative du cycle de vie de sa demande, via les canaux configurés dans ses préférences (in-app, SMS, email). Ces notifications assurent une transparence totale et une expérience client fluide, en informant le client en temps réel sans qu'il ait à consulter activement l'application.")
h3(doc, "2. Règles de Fonctionnement")
bul(doc, "Envoi automatique déclenché par chaque transition de statut ou événement système")
bul(doc, "Canal prioritaire : SMS. Canal secondaire : email. Canal complémentaire : notification in-app")
bul(doc, "En cas d'échec d'envoi : tentative sur canal secondaire + journalisation de l'échec + alerte administrateur")
bul(doc, "Toutes les notifications sont journalisées : type, canal, horodatage, statut de délivrance")
empty(doc)
make_table(doc,
    ["Événement déclencheur", "Contenu de la notification", "Canaux"],
    [
        ["Demande créée avec succès",
         "Votre demande de chéquier n° [REF] a bien été enregistrée le [DATE] à [HEURE]. Un montant de [TOTAL] a été débité de votre compte [COMPTE]. Vous serez notifié à chaque avancement.",
         "In-app, SMS, Email"],
        ["Solde insuffisant (échec)",
         "Votre demande de chéquier n'a pas abouti. Solde insuffisant sur votre compte [COMPTE]. Aucun montant n'a été débité. Veuillez alimenter votre compte et réessayer.",
         "In-app, SMS, Email"],
        ["Service indisponible (échec technique)",
         "Votre demande de chéquier n'a pas pu être enregistrée en raison d'une indisponibilité technique. Aucun montant n'a été débité. Veuillez réessayer ultérieurement.",
         "In-app, SMS"],
        ["Prise en charge (→ En cours)",
         "Votre demande de chéquier n° [REF] est en cours de traitement par notre équipe à l'agence [AGENCE]. Vous serez notifié dès que votre chéquier sera prêt.",
         "In-app, SMS"],
        ["Chéquier imprimé (→ Imprimé / Prêt)",
         "Votre chéquier est prêt ! Présentez-vous à l'agence [NOM AGENCE] — [ADRESSE] muni de votre pièce d'identité pour le retirer. Réf. demande : [REF].",
         "In-app, SMS, Email"],
        ["Chéquier remis (→ Livré)",
         "Votre chéquier (n° [REF]) vous a été remis le [DATE]. Merci pour votre confiance. Pour toute question, contactez votre agence AFBSS.",
         "In-app, SMS"],
    ],
    [4.0, 9.0, 3.2]
)

# ════════════════════════════════════════════════════════════════════════
# VI. MODÈLE DE DONNÉES
# ════════════════════════════════════════════════════════════════════════
h1(doc, "VI. Modèle de Données")

h2(doc, "Entités principales et cardinalités")
empty(doc)
make_table(doc,
    ["Entité 1", "Card. 1", "Entité 2", "Card. 2", "Description / Règle métier clé"],
    [
        ["Client", "1", "Compte bancaire", "N", "Un client peut être titulaire de plusieurs comptes bancaires. Seuls les comptes actifs, non bloqués et disposant de l'autorisation chéquier sont éligibles (BR-01)."],
        ["Compte bancaire", "1", "Demande chéquier", "N", "Un compte peut générer plusieurs demandes. Une demande est obligatoirement rattachée à un seul compte actif et éligible au moment de la soumission."],
        ["Client", "1", "Demande chéquier", "N", "Un client peut soumettre plusieurs demandes. Chaque demande est liée à un et un seul client identifié."],
        ["Compte bancaire", "N", "Agence domiciliation (Home Branch)", "1", "Chaque compte est rattaché à une agence de domiciliation unique. Cette agence reçoit obligatoirement le crédit de commission et taxes, indépendamment de l'agence de livraison choisie (BR-04)."],
        ["Demande chéquier", "1", "Transaction Core Banking", "1", "Chaque demande validée est associée à exactement une transaction atomique dans ALTBank. Impossible de créer une demande dans l'IB sans transaction Core Banking confirmée (BR-02, BR-08)."],
        ["Demande chéquier", "1", "Historique de statuts", "N", "Une demande passe par plusieurs statuts. Chaque transition génère un enregistrement horodaté avec l'identifiant de l'acteur ayant effectué la transition (BR-07)."],
        ["Demande chéquier", "1", "Notification", "N", "Une demande peut générer plusieurs notifications client au cours de son cycle de vie (création, chaque transition de statut, erreur)."],
        ["Demande chéquier", "N", "Agence de livraison", "1", "Le client choisit librement une agence de livraison parmi toutes les agences AFBSS actives (BR-05). L'agence détermine la file d'attente back-office, mais pas l'affectation comptable."],
        ["Opérateur agence", "1", "Changement de statut", "N", "Chaque changement de statut initié par un opérateur est tracé avec son identifiant nominatif. Un opérateur ne peut agir que sur les demandes de son agence (BR-09)."],
        ["Agence AFBSS", "1", "Opérateur agence", "N", "Une agence compte un ou plusieurs opérateurs. Chaque opérateur est rattaché à une seule agence qui détermine son périmètre de travail dans le back-office."],
        ["Session formulaire", "1", "Jeton idempotence", "1", "Chaque instance de formulaire génère un jeton d'idempotence unique. Ce jeton garantit qu'une même demande ne peut être exécutée deux fois (BR-10). Le jeton est invalidé après utilisation ou abandon."],
    ],
    [3.0, 1.0, 3.5, 1.0, 7.7]
)

empty(doc)
h2(doc, "Description des attributs clés des entités principales")

h3(doc, "Entité : Demande chéquier")
make_table(doc,
    ["Attribut", "Type", "Contrainte", "Description"],
    [
        ["Numéro de référence IB", "VARCHAR(20)", "PK, UNIQUE, NOT NULL", "Identifiant unique IB (ex. CHQ-2026-XXXXXX)"],
        ["Référence transaction Core Banking", "VARCHAR(50)", "UNIQUE, NOT NULL", "Lien avec la transaction ALTBank"],
        ["Identifiant client", "VARCHAR(50)", "FK, NOT NULL", "Identifiant du client titulaire"],
        ["Numéro de compte débité", "VARCHAR(30)", "NOT NULL", "Compte source du débit (affiché masqué)"],
        ["Devise", "CHAR(3)", "NOT NULL", "Devise ISO 4217 (USD, SSP, EUR...)"],
        ["Nombre de feuillets", "INTEGER", "NOT NULL, IN (25, 50)", "Nombre de feuillets demandé"],
        ["Montant commission", "DECIMAL(15,2)", "NOT NULL, >= 0", "Commission bancaire débitée"],
        ["Montant taxes", "DECIMAL(15,2)", "NOT NULL, >= 0", "Taxes réglementaires débitées"],
        ["Montant total frais", "DECIMAL(15,2)", "NOT NULL, >= 0", "Total débité (commission + taxes)"],
        ["Code agence de livraison", "VARCHAR(10)", "FK, NOT NULL", "Agence choisie par le client pour le retrait"],
        ["Code agence domiciliation", "VARCHAR(10)", "FK, NOT NULL", "Home Branch — reçoit le crédit des frais (BR-04)"],
        ["Statut courant", "ENUM", "NOT NULL", "EN_ATTENTE / EN_COURS / IMPRIME / LIVRE / ECHOUE"],
        ["Date de création", "TIMESTAMP", "NOT NULL", "Horodatage de création du ticket IB"],
        ["Date de dernière MAJ", "TIMESTAMP", "NOT NULL", "Horodatage de la dernière modification"],
    ],
    [4.2, 2.8, 3.5, 5.7]
)

empty(doc)
h3(doc, "Entité : Historique de statuts")
make_table(doc,
    ["Attribut", "Type", "Contrainte", "Description"],
    [
        ["Identifiant historique", "UUID", "PK, NOT NULL", "Identifiant unique de l'entrée"],
        ["Référence demande", "VARCHAR(20)", "FK, NOT NULL", "Demande concernée"],
        ["Statut source", "ENUM", "NOT NULL", "Statut avant la transition"],
        ["Statut cible", "ENUM", "NOT NULL", "Statut après la transition"],
        ["Identifiant acteur", "VARCHAR(50)", "NOT NULL", "Opérateur ou « SYSTEM »"],
        ["Type acteur", "ENUM", "NOT NULL", "CLIENT / OPERATEUR / SYSTEME"],
        ["Horodatage", "TIMESTAMP", "NOT NULL", "Date et heure précise de la transition"],
        ["Commentaire", "TEXT", "NULLABLE", "Commentaire libre (ex. motif d'un blocage)"],
    ],
    [4.2, 2.8, 3.5, 5.7]
)

empty(doc)
h3(doc, "Entité : Notification")
make_table(doc,
    ["Attribut", "Type", "Contrainte", "Description"],
    [
        ["Identifiant notification", "UUID", "PK, NOT NULL", "Identifiant unique de la notification"],
        ["Référence demande", "VARCHAR(20)", "FK, NOT NULL", "Demande associée"],
        ["Type d'événement", "ENUM", "NOT NULL", "CREATION / ECHEC / PRISE_EN_CHARGE / IMPRIME / LIVRE"],
        ["Canal d'envoi", "ENUM", "NOT NULL", "SMS / EMAIL / IN_APP"],
        ["Destinataire", "VARCHAR(255)", "NOT NULL", "Numéro ou email (affiché masqué)"],
        ["Contenu du message", "TEXT", "NOT NULL", "Corps du message envoyé"],
        ["Statut d'envoi", "ENUM", "NOT NULL", "SUCCES / ECHEC / EN_ATTENTE"],
        ["Horodatage d'envoi", "TIMESTAMP", "NOT NULL", "Date et heure de l'envoi ou de la tentative"],
        ["Code erreur", "VARCHAR(50)", "NULLABLE", "Code d'erreur en cas d'échec d'envoi"],
    ],
    [4.2, 2.8, 3.5, 5.7]
)

# ════════════════════════════════════════════════════════════════════════
doc.save(OUTPUT)
print(f"SUCCESS: {OUTPUT}")
